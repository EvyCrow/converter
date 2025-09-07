import docx as d
from tkinter import *
from tkinter import filedialog
from tkinter import font
from tkinter.messagebox import showwarning, showinfo
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
import customtkinter as ctk
import json
import os


SETTINGS_FILE = "settings.json"


def load_settings():
    """Загружает настройки из JSON-файла."""
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    else:
        return {"theme": "orange.json", "appearance": "Light"}


def save_settings(settings):
    """Сохраняет настройки в JSON-файл."""
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=4)


class App:
    # ctk.set_default_color_theme("orange.json")
    ctk.set_appearance_mode("system")

    def __init__(self, master):
        self.settings = load_settings()
        self.theme_color = self.settings.get("theme", "orange.json")
        ctk.set_default_color_theme(self.theme_color)
        ctk.set_appearance_mode(self.settings.get("appearance"))

        self.master = master
        self.data = []
        self.default_font = ["Times", 12, "bold"]
        self.main_font = self.default_font
        self.fonts = sorted(set([f.split()[0] for f in font.families()]))
        self.name = ""

        # background="#E17C2C"

        self.common_frame = ctk.CTkFrame(master, corner_radius=0)
        self.common_frame.pack(side=TOP, fill=X)
        self.left_frame = ctk.CTkFrame(master, corner_radius=0)
        self.left_frame.pack(side=LEFT, fill=Y)

        # ---------------

        ########################

        self.openButton = ctk.CTkButton(self.common_frame, text="Открыть txt", text_color=("#595959", "#0A0A0A"),
                                        command=self.open_txt, width=15)
        self.openButton.pack(side=LEFT, padx=(5, 0), pady=5, expand=False)

        ########################

        self.wordButton = ctk.CTkButton(self.common_frame, text="Экспорт в docx", text_color=("#595959", "#0A0A0A"),
                                        command=self.word_convert, width=15)
        self.wordButton.pack(side=LEFT, padx=(35, 0), pady=0, expand=False)

        ########################

        self.settingsButton = ctk.CTkButton(self.common_frame, text=".  .  .", text_color=("#595959", "#0A0A0A"),
                                            command=self.open_settings, width=10)
        self.settingsButton.pack(side=RIGHT, expand=False, padx=(0, 10))

        ########################

        self.changeButton = ctk.CTkComboBox(self.left_frame, values=self.fonts, width=120,
                                            state="readonly", command=self.change_font)
        self.changeButton.set(self.default_font[0])
        self.changeButton.pack(anchor="nw", padx=5, pady=(45, 0), expand=False)

        ########################

        self.val = ['8', '9', '10', '11', '12', '14', '16', '18', '20', '22', '24', '26', '28', '36']
        self.changeSize = ctk.CTkComboBox(self.left_frame, values=self.val, width=120, command=self.change_size)
        self.changeSize.set(self.default_font[1])
        self.changeSize.pack(anchor="nw", padx=5, pady=(15, 0), expand=False)

        ########################

        self.name_field = ctk.CTkEntry(self.left_frame, width=120, placeholder_text="Название файла")
        self.name_field.pack(anchor="nw", padx=5, pady=(15, 0), expand=False)

        ########################

        self.table_label = ctk.CTkLabel(self.master, text='Sample', font=(self.main_font[0], self.main_font[1]))
        # self.table_label.configure(
        #    background="#F5F5F5")
        self.table_label.pack(anchor=CENTER, fill=Y, pady=(50, 0))

        ########################

        self.master.title("DOC converter")

    @staticmethod
    def load_data(file_path_):
        main_file = open(file_path_, 'r', newline='', encoding='utf-8')
        datalines = main_file.readlines()
        main_file.close()
        return datalines

    def open_txt(self):
        file_path = filedialog.askopenfilename(
            title="Загрузка .txt",
            filetypes=((".txt файлы", "*.txt"), ("Все файлы", "*.*"))
        )
        if file_path:
            try:
                self.data = self.parser(self.load_data(file_path))
                self.table_label.configure(text=self.data)
                self.name = file_path.split("/")[-1][:-4]
                # print(self.data)
            except Exception as e:
                print(f"Error {e}")

        self.name_field.insert(0, self.name)
        # print(name)

    ##############

    def change_font(self, event=None):
        self.main_font[0] = self.changeButton.get()
        self.table_label.configure(font=(self.main_font[0], self.main_font[1]))
        # print(self.main_font[0])

    def change_size(self, event=None):
        self.main_font[1] = int(self.changeSize.get())
        self.table_label.configure(font=(self.main_font[0], self.main_font[1]))

    ##############

    def open_settings(self):
        root = ctk.CTk()
        root.geometry("200x250")
        root.iconbitmap('logo.ico')
        settings = Settings(root)
        root.mainloop()

    ##############

    def word_convert(self):
        doc = d.Document()

        print(self.data)
        if not self.data:
            showwarning(title="Ошибка", message="Не выбран файл")
            return None
        else:
            # doc.add_paragraph(self.data)
            table = doc.add_table(rows=1, cols=3)
            row = table.rows[0].cells
            row[0].text = '№ п/п'
            row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row[1].text = 'Наименование документа'
            row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row[2].text = 'л. д.'
            row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for id_, name, other in self.data:
                row = table.add_row().cells
                row[0].text = id_
                row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row[1].text = name
                row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row[2].text = other
                row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # I know I could do it earlier,
        # but it's a bit harder than I expected
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font_ = run.font
                        font_.name = self.main_font[0]
                        font_.size = Pt(self.main_font[1])

        paragraph_format = doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = Pt(1.15)
        doc.styles['Normal'].paragraph_format.space_before = 0
        doc.styles['Normal'].paragraph_format.space_after = 0

        savepath = filedialog.askdirectory()
        if savepath:
            doc.save(savepath + '/' + self.name + ".docx")
            showinfo(title="Инфо", message="Файл сохранен")
        else:
            return 0

    ##############

    @staticmethod
    def parser(str_data):
        sep_data = []
        # i = 0
        for str_ in str_data:
            str__ = str_[:-2]
            sep_data.append(str__.split("|"))
        return sep_data


class Settings:
    def __init__(self, master):
        self.master = master
        self.settings = load_settings()

        self.theme = ctk.StringVar(value=ctk.get_appearance_mode())
        self.master.title("Настройки")
        self.switch_button = ctk.CTkSwitch(master, text="Темная тема", command=self.dark_Theme_Switch,
                                           variable=self.theme, onvalue="Dark", offvalue="Light")
        self.switch_button.pack(anchor=W, padx=10, pady=5)

        ############

        self.color_frame = ctk.CTkFrame(master, height=100, width=100)
        self.color_frame.pack(anchor=W, padx=10)

        self.theme_color = ctk.StringVar(value=self.settings.get("theme", "orange.json"))

        self.few_Words = ctk.CTkLabel(self.color_frame, text="Цветовая тема", font=("Ubuntu", 12), height=10)
        self.few_Words.pack(anchor=W, pady=(7, 0), padx=5)
        self.few_Words2 = ctk.CTkLabel(self.color_frame, text="требуется перезапуск", font=("Ubuntu", 8), height=10)
        self.few_Words2.pack(anchor=W, pady=5, padx=5)

        # --

        self.orange = ctk.CTkRadioButton(self.color_frame, text="Апельсин", command=self.change_Color_Theme,
                                         variable=self.theme_color, value="orange.json")
        self.orange.pack(anchor=W, pady=(5, 0), padx=10)

        self.yellow = ctk.CTkRadioButton(self.color_frame, text="Лимон", command=self.change_Color_Theme,
                                         variable=self.theme_color, value="yellow.json")
        self.yellow.pack(anchor=W, pady=(5, 0), padx=10)

        self.carrot = ctk.CTkRadioButton(self.color_frame, text="Морковь", command=self.change_Color_Theme,
                                         variable=self.theme_color, value="carrot.json")
        self.carrot.pack(anchor=W, pady=(5, 0), padx=10)

        self.cherry = ctk.CTkRadioButton(self.color_frame, text="Вишня", command=self.change_Color_Theme,
                                         variable=self.theme_color, value="cherry.json")
        self.cherry.pack(anchor=W, pady=5, padx=10)

    def dark_Theme_Switch(self, event=None):
        ctk.set_appearance_mode(self.theme.get())
        self.settings["appearance"] = self.theme.get()
        save_settings(self.settings)

    def change_Color_Theme(self, event=None):
        self.settings["theme"] = self.theme_color.get()
        save_settings(self.settings)
        # print(self.theme_color.get())


